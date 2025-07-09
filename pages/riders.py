import io, json, html, re
import streamlit as st
from docx import Document                # python‑docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_COLOR_INDEX
from expander import load_abbreviation_dict, expand_abbreviations

def strip_html_tags(text):
    return re.sub(r'</?mark>', '', text)

# ─────────────────────────────────────────────────────────────────────────────
# Helper ─ identify real clause headings only - IMPROVED VERSION
# ─────────────────────────────────────────────────────────────────────────────

# Updated regex to handle the actual document format
HEADER_RE = re.compile(r"^\s*(\d{1,3})\\?[\.\:\-]\s*(.*)$")
# Additional regex to handle existing "Clause X" format with various separators
CLAUSE_RE = re.compile(r"^\s*[Cc]lause\s+(\d{1,3})[\.\:\-]?\s*(.*)$", re.IGNORECASE)

def is_clause_heading(text: str, expected: int | None) -> tuple[bool, int | None]:
    """
    Returns (True,new_expected) if *text* is the next main‑clause heading.
    • Headings must be consecutive numbers (30→31→32 …).
    • Must be reasonable length (≤30 words) and not contain obvious paragraph text
    • Also handles numbered clauses without titles (e.g., "91.")
    • Also handles existing "Clause X" format (e.g., "Clause 31. Trading Limits")
    """
    # Try both regex patterns
    m = HEADER_RE.match(text.strip())
    clause_m = CLAUSE_RE.match(text.strip())
    
    # Use whichever pattern matches
    if clause_m:
        m = clause_m
    elif not m:
        return (False, expected)

    num, title = int(m.group(1)), m.group(2).strip()

    # More flexible consecutive numbering - allow some gaps but prefer consecutive
    if expected is None:
        expected = num          # first heading sets the baseline
    
    # Allow for reasonable clause numbers (don't be too strict about perfect sequence)
    if expected is not None and num < expected:
        return (False, expected)  # Don't go backwards
    
    # Be more lenient with gaps for higher numbered clauses (90+)
    max_gap = 10 if num >= 90 else 5
    if expected is not None and num > expected + max_gap:
        return (False, expected)

    # Handle different cases based on title content
    words = title.split()
    
    # Case 1: No title at all (e.g., "91." or "Clause 91")
    if len(words) == 0:
        return (True, num + 1)
    
    # Case 2: Very short title (1-2 words) - likely legitimate clauses
    if len(words) <= 2:
        # Accept short titles like "Deleted" or legitimate short clause names
        return (True, num + 1)
    
    # Case 3: Long text that looks like paragraph content
    if len(words) > 30:
        return (False, expected)
    
    # Case 4: Text starting with obvious sentence starters - likely paragraph text
    paragraph_starters = [
        'if', 'when', 'should', 'the charterers shall', 'the owners shall',
        'in case', 'notwithstanding', 'subject to', 'provided that',
        'it is understood that', 'all', 'any', 'referring to', 'during the'
    ]
    
    title_lower = title.lower()
    if any(title_lower.startswith(starter) for starter in paragraph_starters):
        return (False, expected)
    
    # Case 5: Text with too many sentence indicators - likely paragraph text
    sentence_indicators = ['shall', 'will', 'must', 'should', 'may', 'can']
    indicator_count = sum(1 for indicator in sentence_indicators if indicator in title_lower)
    if indicator_count > 1:  # More than one indicates sentence text
        return (False, expected)
    
    # Case 6: Accept if it looks like a proper heading
    return (True, num + 1)

def clean_header_text(text: str) -> str:
    """Clean and format header text for display"""
    # Remove escape characters
    text = text.replace('\\', '')
    # Remove leading colons, dashes, and extra spaces
    text = re.sub(r'^[\:\-\s]+', '', text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text.strip())
    return text

# ─────────────────────────────────────────────────────────────────────────────
# Streamlit UI
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="NEMO • Word Formatter", layout="wide")
st.title("🛠 Word Formatter + Abbreviation Expander")

# 1️⃣  Abbreviation dictionary ------------------------------------------------
dict_file = st.sidebar.file_uploader("Upload custom abbreviation dictionary (.xlsx)", type=["xlsx"])
if dict_file:
    st.sidebar.success("Custom dictionary loaded!", icon="✅")
    abbr_dict = load_abbreviation_dict(dict_file)
else:
    # Load the same dictionary as used by the app
    if dict_file:
        abbr_dict = load_abbreviation_dict(dict_file)
    else:
        abbr_dict = load_abbreviation_dict("abbreviations7thJuly.xlsx")
    st.sidebar.info("Using default dictionary")

# 2️⃣  User input -------------------------------------------------------------
raw_text = st.text_area("Paste or type your Word text here 👇",
                        height=400, label_visibility="collapsed",
                        key="riders_input")

chars = len(raw_text)
words = len(raw_text.split()) if raw_text.strip() else 0
st.caption(f"**Characters:** {chars}   **Words:** {words}")

# --------------------------------------------------------------------------- #
# 3️⃣  Expand + format on button click
# --------------------------------------------------------------------------- #
col_go, col_dl = st.columns(2, gap="small")

with col_go:
    go = st.button("🚀 Expand & Format", use_container_width=True)

formatted_bytes = None            # <- will hold .docx
preview_lines   = []              # <- for the HTML preview

if go:
    if not raw_text.strip():
        st.warning("Please enter some text before formatting.")
    else:
        with st.spinner("Expanding abbreviations and creating Word document…"):

            # --- 1. expand abbreviations (plain) ---------------------------
            expanded_plain, expanded_highlighted = expand_abbreviations(raw_text, abbr_dict)

            if not expanded_plain or not expanded_highlighted:
                st.error("Something went wrong during abbreviation expansion")
                st.stop()

            # --- 2. build .docx in memory ----------------------------------
            doc   = Document()
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(10)

            expected_clause_num = None
            
            # Process both plain and highlighted versions
            plain_lines = expanded_plain.splitlines() if expanded_plain else []
            highlighted_lines = expanded_highlighted.splitlines() if expanded_highlighted else []
            
            # Ensure both lists have the same length
            max_len = max(len(plain_lines), len(highlighted_lines))
            while len(plain_lines) < max_len:
                plain_lines.append("")
            while len(highlighted_lines) < max_len:
                highlighted_lines.append("")
            
            for i, (plain_line, highlighted_line) in enumerate(zip(plain_lines, highlighted_lines)):

                # Skip blank lines entirely
                if not plain_line.strip():
                    preview_lines.append("")        # keep blank for preview
                    continue

                # Check if this line is a clause heading (handles both "31." and "Clause 31" formats)
                is_header, expected_clause_num = is_clause_heading(plain_line, expected_clause_num)
                
                if is_header:
                    # Extract number and title using both possible regex patterns
                    m = HEADER_RE.match(plain_line.strip())
                    clause_m = CLAUSE_RE.match(plain_line.strip())
                    
                    # Use whichever pattern matched
                    if clause_m:
                        num = clause_m.group(1)
                        title = clean_header_text(clause_m.group(2))
                        original_format = f"Clause {num}"
                    else:
                        num = m.group(1)
                        title = clean_header_text(m.group(2))
                        # Determine original separator from the line
                        if ':' in plain_line:
                            original_format = f"{num}:"
                        elif '-' in plain_line:
                            original_format = f"{num}-"
                        else:
                            original_format = f"{num}."
                        
                    # Determine if we should include the title or treat it as separate paragraph
                    should_include_title = True
                    remaining_text = ""
                    
                    words = title.split()
                    
                    # Check if title looks like paragraph text that got captured
                    if len(words) > 10:  # Long text - likely paragraph
                        should_include_title = False
                        remaining_text = title
                    elif title.lower().startswith(('in case', 'if', 'referring to', 'during the', 'where and when', 'should the')):
                        should_include_title = False
                        remaining_text = title
                    
                    # Format the clause header consistently (always CAPS, bold, underlined)
                    if not should_include_title or not title or title.lower() == 'deleted':
                        clause_text = f"CLAUSE {num}"
                        if title.lower() == 'deleted':
                            clause_text += ". DELETED"
                    else:
                        clause_text = f"CLAUSE {num}. {title.upper()}"

                    # Check if clause header was changed (standardized)
                    original_line = plain_line.strip()
                    header_was_changed = not original_line.upper().startswith(f"CLAUSE {num}.")
                    
                    # docx – header
                    p = doc.add_paragraph()
                    run = p.add_run(clause_text)
                    run.bold = True
                    run.underline = True
                    
                    # Add highlighting if clause header was standardized
                    if header_was_changed:
                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

                    # paragraph formatting
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # preview - show highlighting if changed
                    if header_was_changed:
                        preview_lines.append(
                            f"<b><u><mark>{html.escape(clause_text)}</mark></u></b>"
                        )
                    else:
                        preview_lines.append(
                            f"<b><u>{html.escape(clause_text)}</u></b>"
                        )
                    
                    # If we have remaining text that should be a separate paragraph, add it
                    if remaining_text:
                        # Check if remaining text has highlights
                        remaining_highlighted = highlighted_line[highlighted_line.find(remaining_text):] if remaining_text in highlighted_line else remaining_text
                        
                        # Add paragraph to docx
                        p = doc.add_paragraph()
                        
                        if '<mark>' in remaining_highlighted:
                            # Process highlighted parts
                            parts = re.split(r'(<mark>.*?</mark>)', remaining_highlighted)
                            for part in parts:
                                if part.startswith('<mark>') and part.endswith('</mark>'):
                                    # This is highlighted text
                                    clean_text = part[6:-7]  # Remove <mark> tags
                                    run = p.add_run(clean_text)
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                else:
                                    # Regular text
                                    p.add_run(part)
                        else:
                            # No highlights in remaining text
                            p.add_run(remaining_text)
                            
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(10)
                        
                        # Set font for all runs
                        for run in p.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                        
                        preview_lines.append(remaining_highlighted)
                else:
                    # regular paragraph - check for abbreviation highlights
                    p = doc.add_paragraph()
                    
                    if '<mark>' in highlighted_line:
                        parts = re.split(r'(<mark>.*?</mark>)', highlighted_line)
                        for part in parts:
                            if part.startswith('<mark>') and part.endswith('</mark>'):
                                clean_text = strip_html_tags(part)
                                run = p.add_run(clean_text)
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            else:
                                p.add_run(strip_html_tags(part))
                    else:
                        # No highlights in this line
                        p.add_run(plain_line)
                    
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(10)
                    
                    # Set font for all runs
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)

                    preview_lines.append(highlighted_line)

            # save the document
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            formatted_bytes = bio.read()

        # --- 3. live preview ---------------------------------------------
        st.subheader("Preview")
        
        # Add highlighting legend
        st.markdown(
            """
            <div style="margin-bottom: 10px; padding: 8px; background: #f0f2f6; border-radius: 4px; font-size: 12px;">
            <strong>🎨 Highlighting Legend:</strong><br>
            <span style="background: #90EE90; padding: 2px 4px; border-radius: 2px;">Green</span> = Clause headers standardized &nbsp;&nbsp;
            <span style="background: #FFFF00; padding: 2px 4px; border-radius: 2px;">Yellow</span> = Abbreviations expanded
            <br><small>💡 Highlights are removable in Word: Select All → Home tab → Text Highlight Color → No Color</small>
            </div>
            """, 
            unsafe_allow_html=True
        )
        
        st.markdown(
            "<div style='background:#fff;border:1px solid #3730a3;"
            "padding:1rem;border-radius:6px;white-space:pre-wrap;"
            "font-family:Arial;font-size:10pt;color:#1e293b;"
            "text-align:justify'>" + "<br>".join(preview_lines) + "</div>",
            unsafe_allow_html=True,
        )

# 4️⃣  Download + copy --------------------------------------------------------
if formatted_bytes:
    with col_dl:
        st.download_button("⬇️ Download .docx",
                           data=formatted_bytes,
                           file_name="formatted_text.docx",
                           mime=("application/vnd.openxmlformats-officedocument."
                                 "wordprocessingml.document"),
                           use_container_width=True)

    escaped = json.dumps(expanded_plain)[1:-1]
    copy_js = f"""
    <script>
    function copyWordText(){{
        navigator.clipboard.writeText("{escaped}").then(
            () => {{
                const btn = document.getElementById("copyWordBtn");
                btn.innerText="✅ Copied!";
                setTimeout(()=>btn.innerText="📋 Copy text",2000);
            }},
            () => alert("Copy failed – please try again."));
    }}
    </script>
    <button id="copyWordBtn"
            onclick="copyWordText()"
            style="background:linear-gradient(135deg,#1e3a8a,#1e40af);color:#fff;
                   border:none;border-radius:6px;padding:0.6rem 1.2rem;
                   font-weight:600;font-size:13px;width:100%;margin-top:0.5rem;
                   box-shadow:0 3px 6px rgba(30,58,138,.3);cursor:pointer;">
        📋 Copy text
    </button>
    """
    st.components.v1.html(copy_js, height=55)